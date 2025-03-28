from datetime import datetime
from sqlalchemy.sql import text
from database.db_manager import DatabaseManager
from database.models import FraudAssessment
import logging

logger = logging.getLogger(__name__)

class FraudeService:
    """Serviço para gerenciar a avaliação de fraudes"""
    
    def __init__(self):
        self.db = DatabaseManager()
    
    def get_processos_suspeitos(self, filters=None):
        """Retorna a lista de processos suspeitos de fraude com suas avaliações"""
        try:
            if filters is None:
                filters = {}
            
            query = """
                SELECT 
                    p.external_id,
                    p.process_number as processo,
                    fa.assessment_date,
                    fa.assessment_result,
                    fa.reason_conclusion
                FROM processes p
                LEFT JOIN agreements a ON a.process_id = p.id
                LEFT JOIN (
                    SELECT external_id, assessment_date, assessment_result, reason_conclusion,
                           ROW_NUMBER() OVER (PARTITION BY external_id ORDER BY assessment_date DESC) as rn
                    FROM fraud_assessments
                ) fa ON fa.external_id = p.external_id AND fa.rn = 1
                WHERE a.fraud_suspect = 1
            """
            
            params = {}
            
            if filters.get('external_id'):
                query += " AND p.external_id = :external_id"
                params['external_id'] = filters['external_id']
            
            if filters.get('processo'):
                query += " AND p.process_number LIKE :processo"
                params['processo'] = f"%{filters['processo']}%"
            
            if filters.get('assessment_result'):
                query += " AND fa.assessment_result = :assessment_result"
                params['assessment_result'] = filters['assessment_result']
            
            if filters.get('reason_conclusion'):
                query += " AND fa.reason_conclusion = :reason_conclusion"
                params['reason_conclusion'] = filters['reason_conclusion']
            
            if filters.get('start_date'):
                query += " AND DATE(fa.assessment_date) >= :start_date"
                params['start_date'] = datetime.strptime(filters['start_date'], '%d/%m/%Y').date()
            
            if filters.get('end_date'):
                query += " AND DATE(fa.assessment_date) <= :end_date"
                params['end_date'] = datetime.strptime(filters['end_date'], '%d/%m/%Y').date()
            
            query += " ORDER BY fa.assessment_date DESC"
            
            session = self.db.Session()
            try:
                result = session.execute(text(query), params)
                rows = result.fetchall()
                
                formatted_results = []
                for row in rows:
                    assessment_date = None
                    if row.assessment_date:
                        assessment_date = datetime.strptime(row.assessment_date.split('.')[0], '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y')
                    
                    formatted_results.append({
                        'external_id': row.external_id,
                        'processo': row.processo,
                        'assessment_date': assessment_date,
                        'assessment_result': row.assessment_result or 'Pendente',
                        'reason_conclusion': row.reason_conclusion
                    })
                
                return formatted_results
            finally:
                session.close()
                
        except Exception as e:
            logger.error(f"Erro ao buscar processos suspeitos: {str(e)}")
            raise
    
    def get_processo_by_id(self, processo_id):
        """
        Busca um processo específico pelo ID
        
        Args:
            processo_id (str): ID externo do processo
            
        Returns:
            dict: Dados do processo ou None se não encontrado
        """
        try:
            session = self.db.Session()
            
            query = """
                SELECT p.*, a.status as acordo_status, a.fraud_suspect as avaliacao,
                       fa.assessment_result, fa.reason_conclusion, fa.assessment_date
                FROM processes p
                LEFT JOIN agreements a ON a.process_id = p.id
                LEFT JOIN (
                    SELECT external_id, assessment_result, reason_conclusion, assessment_date,
                           ROW_NUMBER() OVER (PARTITION BY external_id ORDER BY assessment_date DESC) as rn
                    FROM fraud_assessments
                ) fa ON fa.external_id = p.external_id AND fa.rn = 1
                WHERE p.external_id = :id
            """
            
            result = session.execute(text(query), {'id': processo_id})
            row = result.fetchone()
            
            if not row:
                return None
            
            # Tratar a data com segurança
            data_cadastro = None
            if row.created_at:
                try:
                    if isinstance(row.created_at, str):
                        data_cadastro = datetime.strptime(row.created_at.split('.')[0], '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y')
                    else:
                        data_cadastro = row.created_at.strftime('%d/%m/%Y')
                except Exception as e:
                    logger.warning(f"Erro ao formatar data: {str(e)}")
                    data_cadastro = row.created_at if isinstance(row.created_at, str) else None
                
            return {
                'external_id': row.external_id,
                'processo': row.process_number,
                'data_cadastro': data_cadastro,
                'status': row.status,
                'acordo': row.acordo_status,
                'avaliacao': row.avaliacao,
                'assessment_result': row.assessment_result or 'Pendente',
                'reason_conclusion': row.reason_conclusion,
                'assessment_date': row.assessment_date.strftime('%d/%m/%Y %H:%M') if row.assessment_date else None
            }
            
        except Exception as e:
            logger.error(f"Erro ao buscar processo {processo_id}: {str(e)}")
            raise
        finally:
            session.close()
    
    def save_assessment(self, external_id, data):
        """
        Salva uma avaliação de fraude
        """
        try:
            with self.db.Session() as session:
                # Verifica se já existe uma avaliação para este processo
                assessment = session.query(FraudAssessment).filter_by(external_id=external_id).first()
                
                if not assessment:
                    # Se não existe, cria uma nova
                    assessment = FraudAssessment(
                        external_id=external_id,
                        process_number=data['process_number'],
                        assessment_result=data['assessment_result'],
                        reason_conclusion=data['reason_conclusion'],
                        username=data.get('username')  # Salva o usuário que fez a avaliação
                    )
                    session.add(assessment)
                else:
                    # Se existe, atualiza
                    assessment.assessment_result = data['assessment_result']
                    assessment.reason_conclusion = data['reason_conclusion']
                    assessment.assessment_date = datetime.utcnow()
                    assessment.username = data.get('username')  # Atualiza o usuário que fez a avaliação
                
                session.commit()
                
                return {
                    'external_id': assessment.external_id,
                    'process_number': assessment.process_number,
                    'assessment_result': assessment.assessment_result,
                    'reason_conclusion': assessment.reason_conclusion,
                    'assessment_date': datetime.strptime(assessment.assessment_date.split('.')[0], '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y'),
                    'username': assessment.username
                }
                
        except Exception as e:
            logger.error(f"Erro ao salvar avaliação: {str(e)}")
            raise
    
    def get_historico_avaliacoes(self, external_id):
        """Retorna o histórico de avaliações de um processo"""
        try:
            session = self.db.Session()
            try:
                assessments = session.query(FraudAssessment)\
                    .filter_by(external_id=external_id)\
                    .order_by(FraudAssessment.assessment_date.desc())\
                    .all()
                
                return [{
                    'external_id': assessment.external_id,
                    'assessment_result': assessment.assessment_result,
                    'reason_conclusion': assessment.reason_conclusion,
                    'assessment_date': assessment.assessment_date.strftime('%d/%m/%Y %H:%M:%S')
                } for assessment in assessments]
            finally:
                session.close()
                
        except Exception as e:
            logger.error(f"Erro ao buscar histórico de avaliações: {str(e)}")
            raise
    
    def export_data(self, format='excel'):
        """
        Exporta os dados da avaliação de fraude para diferentes formatos
        """
        try:
            # Busca todos os dados
            data = self.get_processos_suspeitos()
            
            if not data:
                return None
                
            # Prepara os headers
            headers = [
                'ID', 'Processo', 'Data da Avaliação', 
                'Resultado da Avaliação', 'Motivo da Conclusão'
            ]
            
            # Prepara os dados
            rows = []
            for item in data:
                row = [
                    item['external_id'],
                    item['processo'],
                    item['assessment_date'],
                    item['assessment_result'] or '',
                    item['reason_conclusion'] or ''
                ]
                rows.append(row)
            
            if format == 'excel':
                return self._export_to_excel(headers, rows)
            elif format == 'csv':
                return self._export_to_csv(headers, rows)
            elif format == 'word':
                return self._export_to_word(headers, rows)
            else:
                raise ValueError(f'Formato de exportação inválido: {format}')
                
        except Exception as e:
            print(f'Erro ao exportar dados: {str(e)}')
            raise
            
    def _export_to_excel(self, headers, rows):
        """Exporta dados para Excel"""
        import pandas as pd
        from io import BytesIO
        
        df = pd.DataFrame(rows, columns=headers)
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, sheet_name='Avaliações de Fraude', index=False)
            workbook = writer.book
            worksheet = writer.sheets['Avaliações de Fraude']
            
            # Formata o cabeçalho
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'top',
                'bg_color': '#D9D9D9',
                'border': 1
            })
            
            # Aplica o formato ao cabeçalho
            for col_num, value in enumerate(headers):
                worksheet.write(0, col_num, value, header_format)
                
            # Ajusta a largura das colunas
            for i, col in enumerate(df.columns):
                column_len = max(df[col].astype(str).apply(len).max(),
                               len(str(col)) + 2)
                worksheet.set_column(i, i, column_len)
        
        output.seek(0)
        return output.getvalue()
        
    def _export_to_csv(self, headers, rows):
        """Exporta dados para CSV"""
        import csv
        from io import StringIO
        
        output = StringIO()
        writer = csv.writer(output, quoting=csv.QUOTE_ALL)
        
        writer.writerow(headers)
        writer.writerows(rows)
        
        return output.getvalue().encode('utf-8')
        
    def _export_to_word(self, headers, rows):
        """Exporta dados para Word"""
        from docx import Document
        from docx.shared import Inches
        from io import BytesIO
        
        doc = Document()
        doc.add_heading('Relatório de Avaliações de Fraude', 0)
        
        # Adiciona a tabela
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Adiciona o cabeçalho
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            
        # Adiciona as linhas de dados
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)
                
        # Ajusta a largura das colunas
        for cell in table.columns[0].cells:
            cell.width = Inches(1.0)
            
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    def exportar_dados(self, formato):
        """Exporta dados para o formato especificado"""
        return self.export_data(formato)
